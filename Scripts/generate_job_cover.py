from docx import Document
from datetime import datetime
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from tally import read_tally
import os
import platform

def generate_job_cover(data, countryDropdown, typeDropdown):
    
    client_job_info = data['client_job_info']
    handle_log_info = data['handle_log_info']
    
    # Create a new Document
    doc = Document()
    
    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # TITLE
    title = doc.add_paragraph()
    run = title.add_run(generate_title(countryDropdown, typeDropdown))
    run.font.size = Pt(40)
    run.font.name = 'Times New Roman'
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Create a table with two columns for alignment
    table = doc.add_table(rows=0, cols=2)
    
    # Set the width of the first column (labels) to be narrower
    table.columns[0].width = Inches(2)  # Adjust the width as needed
    table.columns[1].width = Inches(5)  # Adjust the width of the value column as needed

    # Function to add a row to the table
    def add_row(label, value):
        row_cells = table.add_row().cells
        label_cell = row_cells[0].paragraphs[0]
        label_run = label_cell.add_run(label)
        label_run.font.size = Pt(16)
        label_run.font.name = 'Times New Roman'
        label_cell.paragraph_format.space_after = Pt(0)  # Reduce space after paragraph

        value_cell = row_cells[1].paragraphs[0]
        value_run = value_cell.add_run(value)
        value_run.font.size = Pt(16)
        value_run.font.name = 'Times New Roman'
        value_cell.paragraph_format.space_after = Pt(0)  # Reduce space after paragraph

    # CLIENT and ADDRESS
    add_row('CLIENT', str(client_job_info['client']))
    add_row('ADDRESS', str(client_job_info['address']))
    add_row('', '') 

    # JOB LOCATION and PROJECT #
    add_row('JOB LOCATION', str(client_job_info['project']))
    add_row('PROJECT #', str(client_job_info['job_#']))
    add_row('', '') 

    # CONTACT and TITLE
    add_row('CONTACT', str(client_job_info['contact']))
    add_row('TITLE', 'Project Manager')
    add_row('', '') 

    # START DATE, PHONE, and FAX
    add_row('START DATE', str(handle_log_info['reviewed_by']['date_time']))
    add_row('PHONE', str(client_job_info['phone']))

    # Save the document
    title_ = generate_title(countryDropdown, typeDropdown)
        
    if platform.system() == 'Windows':
        desktop_path = f'{os.path.expanduser("~")}/OneDrive/Desktop/'
    elif platform.system() == 'Darwin':
        desktop_path = f'{os.path.expanduser("~")}/Desktop/'
    
    job_num = client_job_info['job_#']
    project = get_before_comma(client_job_info['project'])
    folder_name = f'{job_num} - {project}'
    folder_path = f'{desktop_path}{folder_name}'
    
    os.makedirs(folder_path, exist_ok=True)
    output_path = f'{folder_path}/Job Cover {title_}.docx'
    doc.save(output_path)
    print(f"Job cover document saved to {output_path}")
    

def generate_title(countryDropdown, typeDropdown):
    title = ''
    
    if countryDropdown == 'Canada':
        country = 'C'
    else:
        country = 'F'
        
    type = str(typeDropdown)
    tally = read_tally()
    
    now = datetime.now()
    current_year = str(now.year)[-2:]
    current_month = str(now.month).zfill(2)
    
    title += country + type + current_year + current_month + tally
    
    return title

def get_before_comma(input_string):
    # Check if a comma exists in the string
    if ',' in input_string:
        # Split the string at the first comma and return the part before the comma
        return input_string.split(',', 1)[0].strip()
    else:
        # If no comma, return the entire string
        return input_string.strip()