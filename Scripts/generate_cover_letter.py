from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from datetime import date
from tally import read_tally
import re
import os
import sys
import platform

def get_asset_path(filename):
    if getattr(sys, 'frozen', False):
        # When running as a packaged executable
        base_path = sys._MEIPASS
    else:
        return f'../{filename}'
    
    return os.path.join(base_path, filename)

def generate_cover_letter(data, country, type):
    
    client_job_info = data['client_job_info']
    sample_info = data['sample_info']
    handle_log_info = data['handle_log_info']
    
    # Create a new Document
    doc = Document()
    
    style = doc.styles['Normal']
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing = 1.0  # Set line spacing to 1.0
    
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    
    
    if country == 'United States':
        doc.add_picture(get_asset_path('Assets/BSILogoUS.png'), width=Inches(2.4), height=Inches(0.89))
    else:
        doc.add_picture(get_asset_path('Assets/BSILogoCAN.png'), width=Inches(2.4), height=Inches(0.89))
        
    # Header Information
    header = doc.add_paragraph()
    
    run = header.add_run("\n" + str(handle_log_info['reviewed_by']['date_time']) + "\n\n")
    run.font.size = Pt(12)
    run.font.name = 'Calibri'
    
    run = header.add_run(client_job_info['contact'] + "\n")
    run.font.size = Pt(12)
    run.font.name = 'Calibri'
    
    run = header.add_run(client_job_info['client'] + "\n")
    run.font.size = Pt(12)
    run.font.name = 'Calibri'
    
    run = header.add_run((client_job_info['address']) + "\n")
    run.font.size = Pt(12)
    run.font.name = 'Calibri'

    # Subject Line
    subject = doc.add_paragraph()
    subject_run = subject.add_run("Re:      " + content_title(type) + " for " + str(client_job_info['project']) + "\n" + client_job_info['client'] + " Project #" + str(client_job_info['job_#']) +"\n")
    subject_run.bold = True
    subject_run.font.size = Pt(12)
    subject_run.font.name = 'Calibri'
    subject.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Greeting
    greeting = doc.add_paragraph("Dear " + get_title_and_last_name(client_job_info['contact']) + ":")
    greeting_run = greeting.runs[0]
    greeting_run.font.size = Pt(12)
    greeting_run.font.name = 'Calibri'

    # Body Paragraph
    body = doc.add_paragraph()
    body_run = body.add_run(
        "On " + remove_day_of_week(str(handle_log_info['reviewed_by']['date_time'])) + " Building Science Investigations, " + llc_or_inc(country)["name"] + " (\"BSI\") "
        "received " + generate_sample_summary(sample_info) + " from " + client_job_info['client'] + " for the analysis of "
        "fungal growth and spore deposition. The samples were analyzed on " + remove_day_of_week(str(handle_log_info['reviewed_by']['date_time'])) + 
        ". The analytical results are presented on the following pages.\n\n The analytical results and information " 
        "provided in this document are submitted pursuant to BSI’s current terms "
        "and condition of sales including the company’s standard warranty and limitation of liability provisions. No "
        "responsibility is assumed for the manner in which this information is used or interpreted. BSI is not able to "
        "assess any potential health hazard resulting from materials analyzed. This report applies only to the samples "
        "submitted and analyzed. All samples were received in acceptable condition unless otherwise noted. BSI will not "
        "release partial or full copies of this report to any third party without written consent from the client. If you "
        "have any questions please contact me at 716-628-4618."
    )
    body_run.font.size = Pt(12)
    body_run.font.name = 'Calibri'
    body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    
    # Signature Section
    signature = doc.add_paragraph()
    signature_run = signature.add_run("Sincerely,\nBuilding Science Investigations, " + llc_or_inc(country)["name"] + "\n")
    signature_run.font.size = Pt(12)
    signature_run.font.name = 'Calibri'

    # Insert the signature image inline
    inline_image = signature.add_run()
    inline_image.add_picture(get_asset_path('Assets/JonSolomonSignature.png'), width=Inches(1.3), height=Inches(1.15))


    signature_run = signature.add_run("\nJonathan Solomon, Principal\nStructural Mycologist – AIHA EMPAT #158773")
    signature_run.font.size = Pt(12)
    signature_run.font.name = 'Calibri'

    # Footer Information
    footer_section = doc.sections[0]
    footer = footer_section.footer
    footer_paragraph = footer.paragraphs[0]
    footer_run = footer_paragraph.add_run(
        "© Building Science Investigations " + llc_or_inc(country)["abr"] + ", 2024 All Rights Reserved.\n"
        "310 Dolphin Shores Circle Nokomis Florida United States 34275\n"
        "4269 Niagara Boulevard Stevensville Ontario Canada L0S 1S0"
    )
    footer_run.font.size = Pt(10)
    footer_run.font.name = 'Calibri'
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Save the document
    
    if platform.system() == 'Windows':
        desktop_path = f'{os.path.expanduser("~")}/OneDrive/Desktop/'
    elif platform.system() == 'Darwin':
        desktop_path = f'{os.path.expanduser("~")}/Desktop/'
    
    job_num = client_job_info['job_#']
    project = get_before_comma(client_job_info['project'])
    folder_name = f'{job_num} - {project}'
    folder_path = f'{desktop_path}{folder_name}'
    
    os.makedirs(folder_path, exist_ok=True)
    output_path = f'{folder_path}/Cover Letter - {project}.docx'
    doc.save(output_path)
    print(f"Cover letter document saved to {output_path}")
    
def generate_sample_summary(samples):
    matrix_code_dict = {
        "A": "Anderson Air",
        "BA": "Burkard Air",
        "BAF": "Burkard Air Fire-Related Particulate",
        "AOC": "Air-O-Cell",
        "RCS": "RCS Air",
        "B": "Bulk",
        "BF": "Bulk Fire-Related Particulate",
        "S": "Surface",
        "SF": "Surface Fire-Related Particulate",
        "OT": "Other Type"
    }

    # Number to word mapping for numbers < 10
    number_to_word = {
        1: "one", 2: "two", 3: "three", 4: "four", 5: "five",
        6: "six", 7: "seven", 8: "eight", 9: "nine"
    }

    matrix_counts = {}
    
    # Count the occurrences of each sample type
    for sample in samples:
        matrix = sample['matrix']
        if matrix in matrix_code_dict:
            matrix_name = matrix_code_dict[matrix]
            matrix_counts[matrix_name] = matrix_counts.get(matrix_name, 0) + 1
    
    # If no samples exist, return a default message or handle it gracefully
    if not matrix_counts:
        return "no samples"
    
    # Sort the matrix counts by their counts in descending order
    sorted_samples = sorted(matrix_counts.items(), key=lambda x: x[1], reverse=True)
    
    # Convert the counts into a list of parts for the final sentence
    parts = []
    for matrix_name, count in sorted_samples:
        # Use word form for counts less than 10
        count_str = number_to_word.get(count, str(count))
        # Handle singular vs plural form
        part = f"{count_str} {matrix_name} sample" if count == 1 else f"{count_str} {matrix_name} samples"
        parts.append(part)
    
    # Handle the conjunctions ("and") and comma-separated list
    if len(parts) > 1:
        sentence = ", ".join(parts[:-1]) + " and " + parts[-1]
    else:
        sentence = parts[0]
    
    return sentence

def remove_day_of_week(date_str):
    result = re.sub(r'^[A-Za-z]+,\s*', '', date_str)
    return result

def content_title(type):
    if type == 'MA':
        type_name = 'Microbiological Analysis'
    else: 
        type_name = 'Particle Sample'
    return type_name

def get_title_and_last_name(full_name):
    # Split the name by spaces
    name_parts = full_name.split()

    # Ensure the name has at least a title and a last name
    if len(name_parts) < 2:
        return full_name

    # Extract the title and the last name
    title = name_parts[0]  # The first part (e.g., "Mr.")
    last_name = name_parts[-1]  # The last part (e.g., "Kanellos")

    # Return the formatted result
    return f"{title} {last_name}"

import re

def llc_or_inc(country):
    if country == "Canada":
        inc = {
            "name": "Incorporated",
            "abr": "Inc."
        }
        return inc
    else:
        llc = {
            "name": "Limited",
            "abr": "LLC."
        }
        return llc

def get_before_comma(input_string):
    # Check if a comma exists in the string
    if ',' in input_string:
        # Split the string at the first comma and return the part before the comma
        return input_string.split(',', 1)[0].strip()
    else:
        # If no comma, return the entire string
        return input_string.strip()